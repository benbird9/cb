# coding:utf-8

from docx import Document
from docx.shared import Inches, Pt


def genExpr():
    import random
    x = random.randint(10, 999)
    y = random.randint(10, 200)
    #z = random.randint(1, 99 - mul_value)
    expr = ''
    if x % 2 == 0:
        expr = str(x) + ' + ' + str(y)
    else:
        if x > y:
            expr = str(x) + ' - ' + str(y)
        else:
            expr = str(y) + ' - ' + str(x)

    return str(expr + ' = ')

document = Document()

for k in range(0, 20):
    document.add_heading('Date: ______   TimeCost: _______  Score: _______', 1)
    document.add_heading('', 1)

    table = document.add_table(rows=0, cols=3)
    for i in range(0, 20):
        row_cells = table.add_row().cells
        for j in range(0, 3):
            row_cells[j].text = genExpr()

    document.add_page_break()


# document.add_page_break()

document.save('demo.docx')