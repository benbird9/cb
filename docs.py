# coding:utf-8

from docx import Document
from docx.shared import Inches


def genExpr():
    import random
    x = random.randint(1, 9)
    y = random.randint(1, 9)
    mul_value = x * y
    z = random.randint(1, 99 - mul_value)
    expr = ''
    if z % 4 == 0:
        expr = str(z) + ' + ' + str(x) + ' x ' + str(y)
    elif z % 4 == 1:
        expr = str(z) + ' + ' + str(mul_value) + ' / ' + str(x)
    elif z % 4 == 2:
        expr = str(z + x * y) + ' - ' + str(x) + ' x ' + str(y)
    elif z % 4 == 3:
        expr = str(z + x * y) + ' - ' + str(mul_value) + ' / ' + str(x)
    return str(expr + ' = ')

document = Document()

for k in range(0, 20):
    document.add_heading('Date: __________         Score:_________', 0)
    table = document.add_table(rows=0, cols=3)
    for i in range(0, 20):
        row_cells = table.add_row().cells
        for j in range(0, 3):
            row_cells[j].text = genExpr()
    document.add_page_break()


# document.add_page_break()

document.save('demo.docx')